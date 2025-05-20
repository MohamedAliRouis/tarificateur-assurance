from flask import Flask, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from datetime import datetime, timezone, timedelta
from docx import Document
from docx2pdf import convert
import os
import threading
import pytz
import re
import glob
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("tarification.log")
    ]
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

# Path constants
TEMPLATE_DIR = "template_docx"
DOC_DIR = "documents"

# Configuration SQLite
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///devis.db'
db = SQLAlchemy(app)

# Définir le modèle de données
class Devis(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    numero_opportunite = db.Column(db.String(100))
    nom_client = db.Column(db.String(100))
    date_creation = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc), 
                                  onupdate=lambda: datetime.now(timezone.utc))
    type_travaux = db.Column(db.String(50))
    cout_ouvrage = db.Column(db.Float)
    presence_existant = db.Column(db.Boolean)
    client_vip = db.Column(db.Boolean)
    garantie = db.Column(db.String(50))
    souhaite_rcmo = db.Column(db.Boolean)
    assurer_intervenants = db.Column(db.Boolean, default=False)
    destination_ouvrage = db.Column(db.String(100))
    adresse_chantier = db.Column(db.String(200))
    description_ouvrage = db.Column(db.Text)
    taux_do = db.Column(db.Float)    
    taux_trc = db.Column(db.Float)
    taux_rcmo = db.Column(db.Float, nullable=True)
    franchise_rcmo = db.Column(db.Float)
    prime_do = db.Column(db.Float)
    prime_trc = db.Column(db.Float)
    prime_rcmo = db.Column(db.Float)
    prime_totale = db.Column(db.Float)
    montant_dm = db.Column(db.Float, nullable=True)
    montant_maintenance_visite = db.Column(db.Float, nullable=True)
    montant_mesures_conservatoires = db.Column(db.Float, nullable=True)
    montant_rcmo = db.Column(db.Float, nullable=True)
    montant_do = db.Column(db.Float, nullable=True)
    montant_trc = db.Column(db.Float, nullable=True)
    franchise_trc = db.Column(db.Float, nullable=True)
    franchise_maintenance = db.Column(db.Float, nullable=True)

# Créer les tables si elles n'existent pas
with app.app_context():
    db.create_all()
    logger.info("Base de données initialisée avec succès !")

def serialize_devis(devis):
    return {
        'id': devis.id,
        'numero_opportunite': devis.numero_opportunite,
        'nom_client': devis.nom_client,
        'date_creation': devis.date_creation.isoformat(),
        'type_travaux': devis.type_travaux,
        'cout_ouvrage': devis.cout_ouvrage,
        'presence_existant': devis.presence_existant,
        'client_vip': devis.client_vip,
        'garantie': devis.garantie,
        'souhaite_rcmo': devis.souhaite_rcmo,
        'assurer_intervenants': devis.assurer_intervenants,
        'destination_ouvrage': devis.destination_ouvrage,
        'adresse_chantier': devis.adresse_chantier,
        'description_ouvrage': devis.description_ouvrage,
        'taux_do': devis.taux_do,
        'taux_trc': devis.taux_trc,
        'taux_rcmo': devis.taux_rcmo,
        'franchise_rcmo': devis.franchise_rcmo,
        'prime_do': devis.prime_do,
        'prime_trc': devis.prime_trc,
        'prime_rcmo': devis.prime_rcmo,
        'prime_totale': devis.prime_totale,
        'montant_dm': devis.montant_dm,
        'montant_maintenance_visite': devis.montant_maintenance_visite,
        'montant_mesures_conservatoires': devis.montant_mesures_conservatoires,
        'montant_rcmo': devis.montant_rcmo,
        'montant_do': devis.montant_do,
        'montant_trc': devis.montant_trc,
        'franchise_trc': devis.franchise_trc,
        'franchise_maintenance': devis.franchise_maintenance
    }

def sanitize_filename(value):
    return re.sub(r'[^\w\-_. ]', '_', str(value))

def generate_filename(numero_opportunite, dt, ext="docx"):
    safe_numero = sanitize_filename(numero_opportunite)
    date_part = dt.strftime('%d-%m-%Y_%H-%M')
    return f"Proposition_commerciale_{safe_numero}_{date_part}.{ext}"

def to_paris_time(dt):
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(pytz.timezone("Europe/Paris"))

def replace_in_runs(runs, values):
    """Replace placeholder keys in runs with their corresponding values."""
    for run in runs:
        for key, val in values.items():
            if key in run.text:
                run.text = run.text.replace(key, str(val))

def replace_placeholders_in_doc(doc, values):
    """Replace all placeholders in paragraphs and tables in a Word document."""
    # Replace in paragraphs
    for p in doc.paragraphs:
        inline_text = ''.join(run.text for run in p.runs)
        for key, val in values.items():
            if key in inline_text:
                inline_text = inline_text.replace(key, str(val))
        # Remettre le texte fusionné dans un seul run
        if p.runs:
            p.runs[0].text = inline_text
            for i in range(1, len(p.runs)):
                p.runs[i].text = ''

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    inline_text = ''.join(run.text for run in paragraph.runs)
                    for key, val in values.items():
                        if key in inline_text:
                            inline_text = inline_text.replace(key, str(val))
                    if paragraph.runs:
                        paragraph.runs[0].text = inline_text
                        for i in range(1, len(paragraph.runs)):
                            paragraph.runs[i].text = ''

# Fonction pour garantir l'existence du PDF (création si nécessaire)
def ensure_pdf_exists(docx_path):
    """Checks if PDF exists for the given DOCX file, creates it if not.
    Returns the path to the PDF file."""
    pdf_path = docx_path.replace('.docx', '.pdf')
    if not os.path.exists(pdf_path):
        import pythoncom
        pythoncom.CoInitialize()
        try:
            convert(docx_path, pdf_path)
            logger.info(f"PDF généré avec succès : {pdf_path}")
        except Exception as e:
            logger.error(f"Erreur lors de la génération du PDF: {str(e)}")
            raise e
        finally:
            pythoncom.CoUninitialize()
    return pdf_path

# Fonction pour générer le PDF en arrière-plan
def generate_pdf_async(docx_path, pdf_path):
    import pythoncom
    pythoncom.CoInitialize()
    
    try:
        convert(docx_path, pdf_path)
        logger.info(f"PDF généré avec succès en arrière-plan : {pdf_path}")
    except Exception as e:
        logger.error(f"Erreur lors de la génération du PDF en arrière-plan: {str(e)}")
    finally:
        pythoncom.CoUninitialize()


# Fonction pour générer le fichier DOCX en arrière-plan
def generate_docx_file(devis_id):
    devis = Devis.query.get_or_404(devis_id)
    
    # Sélection du template en fonction des options choisies
    template_filename = ""
    
    # Déterminer le template en fonction des garanties choisies
    if devis.garantie == "DO+TRC":
        if devis.souhaite_rcmo:
            template_filename = "template_do_trc_rcmo.docx"
        else:
            template_filename = "template_do_trc.docx"
    elif devis.garantie == "TRC":
        if devis.souhaite_rcmo:
            template_filename = "template_trc_rcmo.docx"
        else:
            template_filename = "template_trc.docx"
    elif devis.garantie == "DO":
        if devis.souhaite_rcmo:
            template_filename = "template_do_rcmo.docx"
        else:
            template_filename = "template_do.docx"
    else:
        # Fallback au cas où la garantie n'est pas reconnue
        template_filename = "template_do.docx"
        logger.warning(f"Garantie non reconnue pour devis {devis_id}: {devis.garantie}. Utilisation du template par défaut.")
    
    template_path = os.path.join(TEMPLATE_DIR, template_filename)
    
    # Vérifier si le template existe
    if not os.path.exists(template_path):
        logger.error(f"Template {template_filename} introuvable. Chemin: {template_path}")
        raise FileNotFoundError(f"Template {template_filename} introuvable.")
    
    doc = Document(template_path)
    
    # Préparer toutes les valeurs dans un seul dictionnaire, sans conditions
    values = {
        # Valeurs communes
        "numero_opportunite": devis.numero_opportunite or "",
        "nom_client": devis.nom_client or "",
        "destination_ouvrage": devis.destination_ouvrage or "",  
        "type_travaux": devis.type_travaux or "",
        "cout_ouvrage": f"{devis.cout_ouvrage:,.2f} €" if devis.cout_ouvrage is not None else "0,00 €",
        "presence_existant": "Oui" if devis.presence_existant else "Non",
        "garantie": devis.garantie or "",
        "description_ouvrage": devis.description_ouvrage or "",
        "adresse_chantier": devis.adresse_chantier or "",
        "date_creation": devis.date_creation.strftime("%d/%m/%Y"),
        
        # Valeurs DO
        "prime_do": f"{devis.prime_do:,.2f} €" if devis.prime_do is not None else "0,00 €",
        "montant_do": f"{devis.montant_do:,.2f} €" if devis.montant_do is not None else "0,00 €",
        
        # Valeurs TRC
        "prime_trc": f"{devis.prime_trc:,.2f} €" if devis.prime_trc is not None else "0,00 €",
        "franchise_trc": f"{devis.franchise_trc:,.2f} €" if devis.franchise_trc is not None else "0,00 €",
        "franchise_maintenance": f"{devis.franchise_maintenance:,.2f} €" if devis.franchise_maintenance is not None else "0,00 €",
        "montant_dm": f"{devis.montant_dm:,.2f} €" if devis.montant_dm is not None else "0,00 €",
        "montant_maintenance_visite": f"{devis.montant_maintenance_visite:,.2f} €" if devis.montant_maintenance_visite is not None else "0,00 €",
        "montant_mesures_conservatoires": f"{devis.montant_mesures_conservatoires:,.2f} €" if devis.montant_mesures_conservatoires is not None else "0,00 €",
        "montant_trc": f"{devis.montant_trc:,.2f} €" if devis.montant_trc is not None else "0,00 €",
        
        # Valeurs RCMO
        "prime_rcmo": f"{devis.prime_rcmo:,.2f} €" if devis.prime_rcmo is not None else "0,00 €",
        "franchise_rcmo": f"{devis.franchise_rcmo:,.2f} €" if devis.franchise_rcmo is not None else "0,00 €",
        "assurer_intervenants": "AVEC" if devis.assurer_intervenants else "SANS",
        "montant_rcmo": f"{devis.montant_rcmo:,.2f} €" if devis.montant_rcmo is not None else "0,00 €",
        
        # Autres valeurs
        "prime_totale": f"{devis.prime_totale:,.2f} €" if devis.prime_totale is not None else "0,00 €",
    }
    
    # Remplacer tous les placeholders dans le document
    replace_placeholders_in_doc(doc, values)
    
    # Convertir UTC en heure Europe/Paris
    date_local = to_paris_time(devis.date_creation)

    filename = generate_filename(devis.numero_opportunite, date_local, "docx")
    filepath = os.path.join(DOC_DIR, filename)
    os.makedirs(DOC_DIR, exist_ok=True)
    doc.save(filepath)
    
    pdf_filepath = filepath.replace('.docx', '.pdf')
    
    pdf_thread = threading.Thread(
        target=generate_pdf_async,
        args=(filepath, pdf_filepath)
    )
    pdf_thread.daemon = True
    pdf_thread.start()
    
    return filepath


def cleanup_old_files(numero_opportunite):
    """Supprimer les anciens fichiers de documents pour un numero_opportunite donné."""
    pattern = os.path.join(DOC_DIR, f"Proposition_commerciale_{sanitize_filename(numero_opportunite)}_*.docx")
    old_docx_files = glob.glob(pattern)
    
    for old_file in old_docx_files:
        try:
            old_pdf = old_file.replace('.docx', '.pdf')
            if os.path.exists(old_file):
                os.remove(old_file)
                logger.info(f"Ancien fichier DOCX supprimé: {old_file}")
            if os.path.exists(old_pdf):
                os.remove(old_pdf)
                logger.info(f"Ancien fichier PDF supprimé: {old_pdf}")
        except Exception as e:
            logger.error(f"Erreur lors de la suppression de l'ancien fichier {old_file}: {str(e)}")


# Fonction de validation des données
def validate_devis_data(data, is_update=False):
    errors = {}
    
    # Champs obligatoires
    required_fields = ['numero_opportunite', 'nom_client', 'type_travaux', 
                       'cout_ouvrage', 'garantie',
                       'adresse_chantier', 'description_ouvrage']
    
    # En cas de création, tous les champs sont obligatoires
    # En cas de mise à jour, vérifier uniquement les champs fournis
    for field in required_fields:
        if (not is_update and (field not in data or data[field] in [None, ''])) or \
           (is_update and field in data and data[field] in [None, '']):
            errors[field] = f"Le champ '{field}' est requis"
    
    # Validation de types
    if 'cout_ouvrage' in data and data['cout_ouvrage'] not in [None, '']:
        try:
            cout = float(data['cout_ouvrage'])
            if cout <= 0:
                errors['cout_ouvrage'] = "Le coût de l'ouvrage doit être supérieur à zéro"
        except (ValueError, TypeError):
            errors['cout_ouvrage'] = "Le coût de l'ouvrage doit être un nombre valide"
    
    # Validation des taux et primes
    for field in ['taux_do', 'taux_trc', 'taux_rcmo', 'prime_do', 'prime_trc', 'prime_rcmo', 'prime_totale']:
        if field in data and data[field] not in [None, '']:
            try:
                val = float(data[field])
                if val < 0:
                    errors[field] = f"Le champ '{field}' ne peut pas être négatif"
            except (ValueError, TypeError):
                errors[field] = f"Le champ '{field}' doit être un nombre valide"
    
    # Validation pour le champ 'garantie'
    if 'garantie' in data and data['garantie'] not in [None, '']:
        valid_garanties = ['DO', 'TRC', 'DO+TRC']  # Ajouter les valeurs valides
        if data['garantie'] not in valid_garanties:
            errors['garantie'] = "La garantie doit être 'DO' 'TRC' ou 'DO+TRC'"
    
    return errors


def clean_empty_fields(data, fields):
    """Convertir les chaînes vides en None pour certains champs spécifiques dans le dictionnaire de données."""
    for field in fields:
        if field in data and data[field] == '':
            data[field] = None
    return data


# Endpoint GET pour récupérer tous les devis
@app.route('/api/devis', methods=['GET'])
def get_devis():
    devis = Devis.query.all()
    return jsonify([serialize_devis(d) for d in devis])

# Endpoint GET pour récupérer un devis spécifique
@app.route('/api/devis/<int:devis_id>', methods=['GET'])
def get_devis_detail(devis_id):
    devis = Devis.query.get_or_404(devis_id)
    return jsonify(serialize_devis(devis))

# Endpoint POST pour créer un nouveau devis
@app.route('/api/devis', methods=['POST'])
def create_devis():
    data = request.json
    logger.info(f"Données reçues: {data}")
    
    # Validation des données
    validation_errors = validate_devis_data(data)
    if validation_errors:
        logger.info(f"Erreurs de validation: {validation_errors}")
        return jsonify({
            'error': 'Erreurs de validation',
            'details': validation_errors
        }), 400
    
    # Clean empty fields - ajouter franchise_rcmo à la liste des champs à nettoyer
    data = clean_empty_fields(data, ['taux_do', 'taux_trc', 'taux_rcmo', 
                                    'montant_dm', 'montant_maintenance_visite', 
                                    'montant_mesures_conservatoires', 'montant_rcmo', 'montant_do',
                                    'montant_trc', 'franchise_trc', 'franchise_maintenance', 
                                    'franchise_rcmo'])  # <- Ajouté franchise_rcmo ici
    

    
    try:
        devis = Devis(
            numero_opportunite=data.get('numero_opportunite'),
            nom_client=data.get('nom_client'),
            type_travaux=data.get('type_travaux'),
            cout_ouvrage=data.get('cout_ouvrage'),
            presence_existant=data.get('presence_existant', False),
            client_vip=data.get('client_vip', False),
            garantie=data.get('garantie'),
            souhaite_rcmo=data.get('souhaite_rcmo', False),
            assurer_intervenants=data.get('assurer_intervenants', False),
            destination_ouvrage=data.get('destination_ouvrage'),
            adresse_chantier=data.get('adresse_chantier'),
            description_ouvrage=data.get('description_ouvrage'),
            taux_do=data.get('taux_do'),
            taux_trc=data.get('taux_trc'),
            taux_rcmo=data.get('taux_rcmo'),
            franchise_rcmo=data.get('franchise_rcmo', 0),
            prime_do=data.get('prime_do', 0),
            prime_trc=data.get('prime_trc', 0),
            prime_rcmo=data.get('prime_rcmo', 0),
            prime_totale=data.get('prime_totale', 0),
            # Add the new fields
            montant_dm=data.get('montant_dm'),
            montant_maintenance_visite=data.get('montant_maintenance_visite'),
            montant_mesures_conservatoires=data.get('montant_mesures_conservatoires'),
            montant_rcmo=data.get('montant_rcmo'),
            montant_do=data.get('montant_do'),
            montant_trc=data.get('montant_trc'),
            franchise_trc=data.get('franchise_trc'),
            franchise_maintenance=data.get('franchise_maintenance')
        )
        db.session.add(devis)
        db.session.commit()
        
        # Nettoyer les anciens fichiers avant de générer le nouveau
        cleanup_old_files(devis.numero_opportunite)
        
        # Générer automatiquement le DOCX après la création
        generate_docx_file(devis.id)
        
        return jsonify({'message': 'Devis créé avec succès', 'id': devis.id}), 201
    except Exception as e:
        db.session.rollback()
        logger.error(f"Erreur lors de la création du devis : {str(e)}")
        return jsonify({'error': f"Erreur lors de la création du devis : {str(e)}"}), 500


# Endpoint PATCH pour modifier un devis existant
@app.route('/api/devis/<int:devis_id>', methods=['PATCH'])
def update_devis(devis_id):
    data = request.json
    devis = Devis.query.get_or_404(devis_id)
    
    # Validation des données pour la mise à jour
    validation_errors = validate_devis_data(data, is_update=True)
    if validation_errors:
        return jsonify({
            'error': 'Erreurs de validation',
            'details': validation_errors
        }), 400
    
    allowed_fields = [
        'nom_client', 'numero_opportunite', 'type_travaux',
        'cout_ouvrage', 'presence_existant', 'client_vip', 'garantie',
        'souhaite_rcmo', 'assurer_intervenants', 'destination_ouvrage', 'adresse_chantier',
        'description_ouvrage', 'taux_do', 'taux_trc', 'taux_rcmo', 'franchise_rcmo',
        'prime_do', 'prime_trc', 'prime_rcmo', 'prime_totale',
        # Add the new fields
        'montant_dm', 'montant_maintenance_visite', 'montant_mesures_conservatoires', 
        'montant_rcmo', 'montant_do', 'montant_trc',
        'franchise_trc', 'franchise_maintenance'
    ]
    
    # Clean empty fields - add the new fields to be cleaned
    data = clean_empty_fields(data, ['taux_do', 'taux_trc', 'taux_rcmo', 
                                    'montant_dm', 'montant_maintenance_visite', 
                                    'montant_mesures_conservatoires', 'montant_rcmo', 'montant_do',
                                    'montant_trc', 'franchise_trc', 'franchise_maintenance',
                                    'franchise_rcmo'])  # Ajouté franchise_rcmo ici
    
    
    try:
        for field in allowed_fields:
            if field in data:
                setattr(devis, field, data[field])
        # Mettre à jour la date de modification
        devis.date_creation = datetime.now(timezone.utc)
        db.session.commit()
        
        # Nettoyer les anciens fichiers avant de générer le nouveau
        cleanup_old_files(devis.numero_opportunite)
        
        # Générer automatiquement le DOCX après la mise à jour
        generate_docx_file(devis.id)
        
        return jsonify({'message': 'Devis mis à jour avec succès'})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Erreur lors de la mise à jour : {str(e)}")
        return jsonify({'error': f"Erreur lors de la mise à jour : {str(e)}"}), 500


# Endpoint GET pour télécharger le DOCX existant
@app.route('/api/devis/<int:devis_id>/docx', methods=['GET'])
def get_docx(devis_id):
    devis = Devis.query.get_or_404(devis_id)
    date_local = to_paris_time(devis.date_creation)
    filename = generate_filename(devis.numero_opportunite, date_local, "docx")
    filepath = os.path.join(DOC_DIR, filename)
    if not os.path.exists(filepath):
        filepath = generate_docx_file(devis_id)
    return send_file(filepath, as_attachment=True)


# Endpoint GET pour générer un fichier PDF pour un devis spécifique
@app.route('/api/devis/<int:devis_id>/pdf', methods=['GET'])
def generate_pdf(devis_id):
    devis = Devis.query.get_or_404(devis_id)
    date_local = to_paris_time(devis.date_creation)
    docx_filename = generate_filename(devis.numero_opportunite, date_local, "docx")
    docx_path = os.path.join(DOC_DIR, docx_filename)
    
    # Générer le DOCX si besoin
    if not os.path.exists(docx_path):
        docx_path = generate_docx_file(devis_id)
    
    try:
        # Utiliser la fonction commune pour garantir l'existence du PDF
        pdf_path = ensure_pdf_exists(docx_path)
        return send_file(pdf_path, as_attachment=True)
    except Exception as e:
        return jsonify({'error': f"Erreur lors de la génération du PDF: {str(e)}"}), 500


if __name__ == '__main__':
    with app.app_context():     
        db.create_all()         
    app.run(debug=True)